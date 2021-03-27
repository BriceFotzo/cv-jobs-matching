from docx import Document

document = Document('cv_blue.docx')

# Dictionary = {‘sea’: “ocean”}

if __name__=="__main__":
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph)